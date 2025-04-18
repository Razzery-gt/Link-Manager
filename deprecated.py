import webbrowser
import json
import os
import getpass
import bcrypt
import validators
import requests
import csv
import yaml
import xml.etree.ElementTree as ET
from docx import Document
import pandas as pd
import tkinter as tk
from tkinter import filedialog
import pyperclip
from datetime import datetime
from colorama import init, Fore
import re  # Импорт модуля для регулярных выражений
import logging  # Импорт модуля для логирования

# Инициализация библиотеки colorama
init(autoreset=True)

# Константы
DOCUMENTS_DIR = os.path.join(os.path.expanduser("~"), "Documents", "LinkManager files")
LINKS_FILENAME = os.path.join(DOCUMENTS_DIR, 'url_links.json')
SETTINGS_FILENAME = os.path.join(DOCUMENTS_DIR, 'settings.json')
LOG_FILENAME = os.path.join(DOCUMENTS_DIR, 'link_manager.log')

# Настройка логирования
logging.basicConfig(filename=LOG_FILENAME, level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Начальные настройки
default_settings = {
    "password": bcrypt.hashpw("1234".encode(), bcrypt.gensalt()).decode(),
    "password_required": False,
    "show_links": True,
    "log_level": "INFO",  # Уровень логирования по умолчанию
    "debug_mode": False  # Режим отладки по умолчанию
}

# Стандартные ссылки
default_links = {
    "открыть_браузер": {"url": "https://example.com", "date_added": str(datetime.now()), "category": "Общее", "description": ""},
    "гугл": {"url": "https://www.google.com", "date_added": str(datetime.now()), "category": "Поиск", "description": ""},
    "яндекс": {"url": "https://www.yandex.ru", "date_added": str(datetime.now()), "category": "Поиск", "description": ""}
}

# Функции
def load_links():
    if os.path.exists(LINKS_FILENAME):
        try:
            with open(LINKS_FILENAME, 'r', encoding='utf-8') as f:
                links = json.load(f)
                # Преобразование старых форматов в новый формат с датой добавления и категорией
                for key, value in links.items():
                    if isinstance(value, str):  # Старый формат, только URL
                        links[key] = {"url": value, "date_added": str(datetime.now()), "category": "Без категории", "description": ""}
                    elif "category" not in value:
                        value["category"] = "Без категории"
                        value["description"] = ""
                    elif "description" not in value:
                        value["description"] = ""
                logging.info("Ссылки загружены из файла.")
                return links
        except (json.JSONDecodeError, IOError) as e:
            print(Fore.RED + f"Ошибка загрузки ссылок: {e}.")
            logging.error(f"Ошибка загрузки ссылок: {e}")
    logging.info("Загружены стандартные ссылки.")
    return default_links.copy()

def save_links(links):
    try:
        with open(LINKS_FILENAME, 'w', encoding='utf-8') as f:
            json.dump(links, f, ensure_ascii=False, indent=4)
        logging.info("Ссылки сохранены в файл.")
    except IOError as e:
        print(Fore.RED + f"Ошибка при сохранении ссылок: {e}.")
        logging.error(f"Ошибка при сохранении ссылок: {e}")

def load_settings():
    if os.path.exists(SETTINGS_FILENAME):
        try:
            with open(SETTINGS_FILENAME, 'r', encoding='utf-8') as f:
                settings = json.load(f)
                # Добавление новых настроек, если они отсутствуют
                for key, value in default_settings.items():
                    if key not in settings:
                        settings[key] = value
                logging.info("Настройки загружены из файла.")
                return settings
        except (json.JSONDecodeError, IOError) as e:
            print(Fore.RED + f"Ошибка загрузки настроек: {e}.")
            logging.error(f"Ошибка загрузки настроек: {e}")
    logging.info("Загружены стандартные настройки.")
    return default_settings.copy()

def save_settings(settings):
    try:
        with open(SETTINGS_FILENAME, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=4)
        logging.info("Настройки сохранены в файл.")
    except IOError as e:
        print(Fore.RED + f"Ошибка при сохранении настроек: {e}.")
        logging.error(f"Ошибка при сохранении настроек: {e}")

def open_browser(url):
    try:
        webbrowser.open(url)
        print(Fore.GREEN + f"Открываем: {url}")
        logging.info(f"Открыта ссылка: {url}")
    except Exception as e:
        print(Fore.RED + f"Ошибка при открытии браузера: {e}")
        logging.error(f"Ошибка при открытии браузера: {e}")

def is_valid_url(url):
    return validators.url(url)

def check_url_accessibility(url):
    try:
        response = requests.get(url, timeout=5)
        if response.status_code == 200:
            logging.info(f"URL '{url}' доступен (статус {response.status_code}).")
            return True
        else:
            logging.warning(f"URL '{url}' вернул статус {response.status_code}.")
            return False
    except requests.RequestException:
        logging.error(f"Ошибка при проверке доступности URL '{url}'.")
        return False

def show_available_keys(links):
    if links:
        print(Fore.CYAN + "Доступные ключи для открытия браузера:")
        for index, data in enumerate(links.values(), 1):
            print(f"{index}. {Fore.YELLOW}{list(links.keys())[index - 1]} - {data['url']} (Категория: {data['category']}, Добавлено: {data['date_added']}, Описание: {data['description']})")
    else:
        print(Fore.RED + "Нет доступных ключей.")

def reset_program():
    if os.path.exists(LINKS_FILENAME):
        os.remove(LINKS_FILENAME)
        logging.warning("Файл со ссылками удален.")
    if os.path.exists(SETTINGS_FILENAME):
        os.remove(SETTINGS_FILENAME)
        logging.warning("Файл с настройками удален.")
    print("Link Manager сброшен к настройкам по умолчанию.")
    logging.info("Программа сброшена к настройкам по умолчанию.")

def hash_password(password):
    hashed = bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()
    logging.info("Пароль захэширован.")
    return hashed

def verify_password(stored_password, provided_password):
    verified = bcrypt.checkpw(provided_password.encode(), stored_password.encode())
    if verified:
        logging.info("Пароль верифицирован.")
    else:
        logging.warning("Неверный пароль.")
    return verified

def count_links(links):
    count = len(links)
    print(Fore.GREEN + f"Количество сохраненных ссылок: {count}")
    logging.info(f"Подсчитано количество ссылок: {count}")

def copy_to_clipboard(text):
    try:
        pyperclip.copy(text)
        print(Fore.GREEN + "Ссылка скопирована в буфер обмена.")
        logging.info(f"Ссылка '{text}' скопирована в буфер обмена.")
    except pyperclip.PyperclipException:
        print(Fore.RED + "Ошибка при копировании в буфер обмена.  Убедитесь, что установлена библиотека pyperclip и работает буфер обмена.")
        logging.error("Ошибка при копировании в буфер обмена.")

def choose_file(save=False, filetypes=(("JSON files", "*.json"), ("All files", "*.*"))):
    root = tk.Tk()
    root.withdraw()  # Скрыть главное окно
    if save:
        filepath = filedialog.asksaveasfilename(initialdir=DOCUMENTS_DIR, title="Выберите место для сохранения", filetypes=filetypes)
        if filepath:
            logging.info(f"Выбран файл для сохранения: {filepath}")
    else:
        filepath = filedialog.askopenfilename(initialdir=DOCUMENTS_DIR, title="Выберите файл для импорта", filetypes=filetypes)
        if filepath:
            logging.info(f"Выбран файл для импорта: {filepath}")
    return filepath

# --- Функции экспорта ---
def export_links(links, filename, format):
    try:
        if format == 'csv':
            export_to_csv(links, filename)
        elif format == 'json':
            export_to_json(links, filename)
        elif format == 'yaml':
            export_to_yaml(links, filename)
        elif format == 'xml':
            export_to_xml(links, filename)
        elif format == 'docx':
            export_to_docx(links, filename)
        elif format == 'txt':
            export_to_txt(links, filename)
        elif format == 'xlsx':
            export_to_xlsx(links, filename)
        else:
            print(Fore.RED + "Неподдерживаемый формат файла.")
            logging.warning(f"Попытка экспорта в неподдерживаемом формате: {format}")
            return

        print(Fore.GREEN + f"Ссылки экспортированы в {filename} в формате {format.upper()}.")
        logging.info(f"Ссылки экспортированы в {filename} в формате {format.upper()}.")

    except Exception as e:
        print(Fore.RED + f"Ошибка при экспорте ссылок: {e}.")
        logging.error(f"Ошибка при экспорте ссылок: {e}")

def export_to_csv(links, filename):
    with open(filename, mode='w', newline='', encoding='utf-8') as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(['Key', 'URL', 'Date Added', 'Category', 'Description'])
        for key, data in links.items():
            writer.writerow([key, data['url'], data['date_added'], data['category'], data['description']])
    logging.info(f"Ссылки экспортированы в CSV: {filename}")

def export_to_json(links, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(links, f, ensure_ascii=False, indent=4)
    logging.info(f"Ссылки экспортированы в JSON: {filename}")

def export_to_yaml(links, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        yaml.dump(links, f, allow_unicode=True, indent=4, stream=f)
    logging.info(f"Ссылки экспортированы в YAML: {filename}")

def export_to_xml(links, filename):
    root = ET.Element("links")
    for key, data in links.items():
        link = ET.SubElement(root, "link")
        ET.SubElement(link, "key").text = key
        ET.SubElement(link, "url").text = data['url']
        ET.SubElement(link, "date_added").text = data['date_added']
        ET.SubElement(link, "category").text = data['category']
        ET.SubElement(link, "description").text = data['description']

    tree = ET.ElementTree(root)
    tree.write(filename, encoding='utf-8', xml_declaration=True)
    logging.info(f"Ссылки экспортированы в XML: {filename}")

def export_to_docx(links, filename):
    document = Document()
    document.add_heading('Links', level=1)
    for key, data in links.items():
        document.add_paragraph(f"Key: {key}")
        document.add_paragraph(f"URL: {data['url']}")
        document.add_paragraph(f"Date Added: {data['date_added']}")
        document.add_paragraph(f"Category: {data['category']}")
        document.add_paragraph(f"Description: {data['description']}")
        document.add_paragraph()

    document.save(filename)
    logging.info(f"Ссылки экспортированы в DOCX: {filename}")

def export_to_txt(links, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        for key, data in links.items():
            f.write(f"Key: {key}\n")
            f.write(f"URL: {data['url']}\n")
            f.write(f"Date Added: {data['date_added']}\n")
            f.write(f"Category: {data['category']}\n")
            f.write(f"Description: {data['description']}\n\n")
    logging.info(f"Ссылки экспортированы в TXT: {filename}")

def export_to_xlsx(links, filename):
    data = []
    for key, link_data in links.items():
        data.append([key, link_data['url'], link_data['date_added'], link_data['category'], link_data['description']])

    df = pd.DataFrame(data, columns=['Key', 'URL', 'Date Added', 'Category', 'Description'])
    df.to_excel(filename, index=False)
    logging.info(f"Ссылки экспортированы в XLSX: {filename}")

# --- Функции импорта ---
def import_links(filename, format):
    try:
        if format == 'csv':
            new_links = import_from_csv(filename)
        elif format == 'json':
            new_links = import_from_json(filename)
        elif format == 'yaml':
            new_links = import_from_yaml(filename)
        elif format == 'xml':
            new_links = import_from_xml(filename)
        elif format == 'docx':
            new_links = import_from_docx(filename)
        elif format == 'txt':
            new_links = import_from_txt(filename)
        elif format == 'xlsx':
            new_links = import_from_xlsx(filename)
        else:
            print(Fore.RED + "Неподдерживаемый формат файла.")
            logging.warning(f"Попытка импорта из неподдерживаемого формата: {format}")
            return

        # Добавляем новые ссылки, проверяя дубликаты
        imported_count = 0
        skipped_duplicates = 0
        for key, data in new_links.items():
            if key in url_links:
                print(Fore.YELLOW + f"Ключ '{key}' уже существует. Пропускаем.")
                logging.warning(f"Пропущен дубликат ключа при импорте: '{key}'")
                skipped_duplicates += 1
            elif data['url'] in [link['url'] for link in url_links.values()]:
                print(Fore.YELLOW + f"Ссылка '{data['url']}' уже существует. Пропускаем.")
                logging.warning(f"Пропущена дублирующаяся ссылка при импорте: '{data['url']}'")
                skipped_duplicates += 1
            else:
                url_links[key] = data
                print(Fore.GREEN + f"Импортирована ссылка: {key} - {data['url']}")
                logging.info(f"Импортирована ссылка: {key} - {data['url']}")
                imported_count += 1

        save_links(url_links)
        print(Fore.GREEN + f"Импортировано {imported_count} ссылок из {filename} в формате {format.upper()}. Пропущено {skipped_duplicates} дубликатов.")
        logging.info(f"Импортировано {imported_count} ссылок из {filename} в формате {format.upper()}. Пропущено {skipped_duplicates} дубликатов.")

    except Exception as e:
        print(Fore.RED + f"Ошибка при импорте ссылок: {e}.")
        logging.error(f"Ошибка при импорте ссылок: {e}")


def import_from_csv(filename):
    links = {}
    with open(filename, mode='r', encoding='utf-8') as csv_file:
        reader = csv.reader(csv_file)
        next(reader, None)  # Skip header
        for row in reader:
            if len(row) == 5:
                key, url, date_added, category, description = row
                if is_valid_url(url):
                    links[key] = {'url': url, 'date_added': date_added, 'category': category, 'description': description}
                else:
                    print(Fore.RED + f"Неверный URL '{url}' в строке '{row}'. Пропускаем.")
                    logging.warning(f"Неверный URL '{url}' в строке CSV '{row}'. Пропущено.")
            elif len(row) == 4: # Обработка файлов без описания для совместимости
                key, url, date_added, category = row
                if is_valid_url(url):
                    links[key] = {'url': url, 'date_added': date_added, 'category': category, 'description': ""}
                else:
                    print(Fore.RED + f"Неверный URL '{url}' в строке '{row}'. Пропускаем.")
                    logging.warning(f"Неверный URL '{url}' в строке CSV '{row}'. Пропущено.")
            else:
                print(Fore.RED + f"Неверное количество столбцов в строке CSV: '{row}'. Пропускаем.")
                logging.warning(f"Неверное количество столбцов в строке CSV: '{row}'. Пропущено.")
    logging.info(f"Ссылки импортированы из CSV: {filename}")
    return links


def import_from_json(filename):
    with open(filename, 'r', encoding='utf-8') as f:
        data = json.load(f)
        for key, value in data.items():
            if isinstance(value, str):
                data[key] = {"url": value, "date_added": str(datetime.now()), "category": "Без категории", "description": ""}
            elif "description" not in value:
                value["description"] = ""
        logging.info(f"Ссылки импортированы из JSON: {filename}")
        return data

def import_from_yaml(filename):
    with open(filename, 'r', encoding='utf-8') as f:
        data = yaml.safe_load(f)
        if data:
            for key, value in data.items():
                if isinstance(value, str):
                    data[key] = {"url": value, "date_added": str(datetime.now()), "category": "Без категории", "description": ""}
                elif "description" not in value:
                    value["description"] = ""
        logging.info(f"Ссылки импортированы из YAML: {filename}")
        return data or {}

def import_from_xml(filename):
    tree = ET.parse(filename)
    root = tree.getroot()
    links = {}
    for link in root.findall('link'):
        key = link.find('key').text
        url = link.find('url').text
        date_added = link.find('date_added').text
        category = link.find('category').text
        description = link.find('description').text if link.find('description') is not None else ""
        links[key] = {'url': url, 'date_added': date_added, 'category': category, 'description': description}
    logging.info(f"Ссылки импортированы из XML: {filename}")
    return links

def import_from_docx(filename):
    document = Document(filename)
    links = {}
    key = None
    url = None
    date_added = None
    category = None
    description = None

    for paragraph in document.paragraphs:
        text = paragraph.text
        if text.startswith("Key: "):
            key = text[5:]
        elif text.startswith("URL: "):
            url = text[5:]
        elif text.startswith("Date Added: "):
            date_added = text[12:]
        elif text.startswith("Category: "):
            category = text[10:]
        elif text.startswith("Description: "):
            description = text[13:]

        if key and url and date_added and category:
            links[key] = {'url': url, 'date_added': date_added, 'category': category, 'description': description or ""}
            key = None
            url = None
            date_added = None
            category = None
            description = None
    logging.info(f"Ссылки импортированы из DOCX: {filename}")
    return links

def import_from_txt(filename):
    links = {}
    key = None
    url = None
    date_added = None
    category = None
    description = None

    with open(filename, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if line.startswith("Key: "):
                key = line[5:]
            elif line.startswith("URL: "):
                url = line[5:]
            elif line.startswith("Date Added: "):
                date_added = line[12:]
            elif line.startswith("Category: "):
                category = line[10:]
            elif line.startswith("Description: "):
                description = line[13:]

            if key and url and date_added and category:
                links[key] = {'url': url, 'date_added': date_added, 'category': category, 'description': description or ""}
                key = None
                url = None
                date_added = None
                category = None
                description = None
    logging.info(f"Ссылки импортированы из TXT: {filename}")
    return links

def import_from_xlsx(filename):
    df = pd.read_excel(filename)
    links = {}
    for index, row in df.iterrows():
        key = str(row['Key'])  # Convert key to string to handle numeric keys
        url = row['URL']
        date_added = str(row['Date Added'])  # Convert date to string
        category = str(row['Category'])
        description = str(row.get('Description', '')) # Получаем описание, если есть
        links[key] = {'url': url, 'date_added': date_added, 'category': category, 'description': description}
    logging.info(f"Ссылки импортированы из XLSX: {filename}")
    return links

def search_links(links, query, search_type='keyword', filters=None):
    found_links = {}

    for key, data in links.items():
        include = False

        if search_type == 'regex':
            if re.search(query, key, re.IGNORECASE) or re.search(query, data['url'], re.IGNORECASE) or re.search(query, data['category'], re.IGNORECASE) or re.search(query, data['description'], re.IGNORECASE):
                include = True
        else:  # search_type == 'keyword'
            if query.lower() in key.lower() or query.lower() in data['url'].lower() or query.lower() in data['category'].lower() or query.lower() in data['description'].lower():
                include = True

        if filters:
            if filters.get('category') and data['category'].lower() != filters['category'].lower():
                include = False
            if filters.get('date_from'):
                try:
                    link_date = datetime.strptime(data['date_added'], '%Y-%m-%d %H:%M:%S.%f')
                    filter_date = datetime.strptime(filters['date_from'], '%Y-%m-%d')
                    if link_date < filter_date:
                        include = False
                except ValueError:
                    print(Fore.RED + "Неверный формат даты в фильтре.")
            if filters.get('date_to'):
                try:
                    link_date = datetime.strptime(data['date_added'], '%Y-%m-%d %H:%M:%S.%f')
                    filter_date = datetime.strptime(filters['date_to'], '%Y-%m-%d')
                    if link_date > filter_date:
                        include = False
                except ValueError:
                    print(Fore.RED + "Неверный формат даты в фильтре.")
            if filters.get('status'):
                accessible = check_url_accessibility(data['url'])
                if (filters['status'] == 'accessible' and not accessible) or (filters['status'] == 'inaccessible' and accessible):
                    include = False

        if include:
            found_links[key] = data

    if found_links:
        print(Fore.CYAN + "Найденные ссылки:")
        for key, data in found_links.items():
            print(f"{Fore.YELLOW}{key}: {data['url']} (Категория: {data['category']}, Добавлено: {data['date_added']}, Описание: {data['description']})")
        logging.info(f"Выполнен поиск. Найдено {len(found_links)} ссылок.")
    else:
        print(Fore.RED + "Ссылки не найдены.")
        logging.info("Поиск не дал результатов.")

def show_statistics(links):
    category_counts = {}
    for data in links.values():
        category = data['category']
        category_counts[category] = category_counts.get(category, 0) + 1

    print(Fore.CYAN + "\nСтатистика:")
    print(Fore.YELLOW + "Количество ссылок по категориям:")
    for category, count in category_counts.items():
        print(f"- {category}: {count}")

    # Дополнительную статистику можно добавить здесь

    logging.info("Показана статистика.")

def set_log_level(settings):
    level_str = settings.get("log_level", "INFO").upper()
    if hasattr(logging, level_str):
        logging.getLogger().setLevel(getattr(logging, level_str))
        print(Fore.GREEN + f"Уровень логирования установлен на: {level_str}")
        logging.info(f"Уровень логирования установлен на: {level_str}")
    else:
        print(Fore.RED + f"Неверный уровень логирования: {level_str}. Установлен уровень INFO.")
        logging.warning(f"Неверный уровень логирования в настройках: {level_str}. Использован уровень INFO.")
        logging.getLogger().setLevel(logging.INFO)
        settings["log_level"] = "INFO"
        save_settings(settings)

def run_debug_functions(links):
    print(Fore.CYAN + "\nОтладочные функции:")
    print("1. Вывести все настройки")
    print("2. Тестировать доступность URL")
    print("3. Очистить файл логов")
    print("4. Назад")

    choice = menu_option("Выберите действие: ", range(1, 5))

    if choice == 1:
        print(Fore.YELLOW + "Текущие настройки:")
        print(json.dumps(settings, indent=4, ensure_ascii=False))
        logging.debug("Выведены текущие настройки.")
    elif choice == 2:
        url_to_test = input("Введите URL для тестирования: ")
        if is_valid_url(url_to_test):
            if check_url_accessibility(url_to_test):
                print(Fore.GREEN + f"URL '{url_to_test}' доступен.")
                logging.debug(f"URL '{url_to_test}' признан доступным через отладочную функцию.")
            else:
                print(Fore.RED + f"URL '{url_to_test}' недоступен.")
                logging.debug(f"URL '{url_to_test}' признан недоступным через отладочную функцию.")
        else:
            print(Fore.RED + "Неверный URL.")
    elif choice == 3:
        try:
            with open(LOG_FILENAME, 'w') as f:
                f.write("")
            print(Fore.GREEN + "Файл логов очищен.")
            logging.info("Файл логов очищен через отладочную функцию.")
        except IOError as e:
            print(Fore.RED + f"Ошибка при очистке файла логов: {e}")
            logging.error(f"Ошибка при очистке файла логов через отладочную функцию: {e}")
    elif choice == 4:
        pass

def menu_option(prompt, options):
    while True:
        try:
            choice = int(input(prompt))
            if choice in options:
                return choice
            else:
                print(Fore.RED + "Неверный ввод. Пожалуйста, попробуйте снова.")
        except ValueError:
            print(Fore.RED + "Неверный ввод. Пожалуйста, попробуйте снова.")

# Основной цикл
os.makedirs(DOCUMENTS_DIR, exist_ok=True)
url_links = load_links()
settings = load_settings()
set_log_level(settings) # Установка уровня логирования при запуске

# Проверка пароля
if settings["password_required"]:
    password_attempts = 3
    while password_attempts > 0:
        password_input = getpass.getpass("Введите пароль для доступа к Link Manager: ")
        if verify_password(settings["password"], password_input):
            break
        else:
            password_attempts -= 1
            print(Fore.RED + f"Неверный пароль. Осталось попыток: {password_attempts}")
            logging.warning(f"Неудачная попытка ввода пароля. Осталось {password_attempts} попыток.")
    else:
        print(Fore.RED + "Доступ запрещен. Нажмите любую клавишу, чтобы продолжить...")
        input()
        logging.critical("Доступ к программе запрещен из-за неверного пароля.")
        exit()

print(Fore.GREEN + "Добро пожаловать в Link Manager!")
print(Fore.GREEN + "Версия: 1.2 (с расширенными функциями)")
logging.info("Программа запущена.")

while True:
    print("\nВыберите действие:")
    print("1. Открыть ссылку по ключу")
    print("2. Добавить новую ссылку или изменить существующую")
    print("3. Удалить ссылку")
    print("4. Переименовать ссылку")
    print("5. Показать доступные ключи")
    print("6. Настройки")
    print("7. Подсчитать количество сохраненных ссылок")
    print("8. Экспорт ссылок")
    print("9. Импорт ссылок")
    print("10. Поиск ссылок")
    print("11. Статистика")
    print("12. Выход")

    choice = menu_option("Введите номер действия: ", range(1, 13))

    if choice == 1:
        if settings["show_links"]:
            show_available_keys(url_links)
        user_input = input("Введите ключ для открытия браузера (или введите номер, или 'copy'): ").strip()

        if user_input.lower() == 'copy':
            key_to_copy = input("Введите ключ ссылки для копирования: ")
            if key_to_copy in url_links:
                copy_to_clipboard(url_links[key_to_copy]['url'])
            else:
                print(Fore.RED + f"Ключ '{key_to_copy}' не найден.")
        elif user_input.isdigit():
            selected_index = int(user_input) - 1
            if 0 <= selected_index < len(url_links):
                key_list = list(url_links.keys())
                selected_key = key_list[selected_index]
                selected_url = url_links[selected_key]['url']
                if check_url_accessibility(selected_url):
                    open_browser(selected_url)
                else:
                    print(Fore.RED + f"Не удалось получить доступ к URL: {selected_url}")
            else:
                print(Fore.RED + "Неверный номер ключа.")
        elif user_input.lower() in [k.lower() for k in url_links.keys()]:
            selected_key = next(k for k in url_links.keys() if k.lower() == user_input.lower())
            selected_url = url_links[selected_key]['url']
            if check_url_accessibility(selected_url):
                open_browser(selected_url)
            else:
                print(Fore.RED + f"Не удалось получить доступ к URL: {selected_url}")
        else:
            print(Fore.RED + f"Ключ '{user_input}' не найден.")

    elif choice == 2:
        new_key = input("Введите ключ: ")
        new_url = input("Введите URL: ")
        new_category = input("Введите категорию: ")
        new_description = input("Введите описание: ")

        if not is_valid_url(new_url):
            print(Fore.RED + "Неверный URL. Пожалуйста, введите корректный URL.")
            continue

        if new_url in [link['url'] for link in url_links.values()]:
            print(Fore.RED + "Эта ссылка уже существует.")
            continue
        url_links[new_key] = {"url": new_url, "date_added": str(datetime.now()), "category": new_category, "description": new_description}
        save_links(url_links)
        print(Fore.GREEN + f"Ссылка для ключа '{new_key}' добавлена/обновлена.")
        logging.info(f"Добавлена/обновлена ссылка: '{new_key}' - '{new_url}' (Категория: '{new_category}', Описание: '{new_description}')")

    elif choice == 3:
        key_to_delete = input("Введите ключ для удаления: ")
        if key_to_delete in url_links:
            del url_links[key_to_delete]
            save_links(url_links)
            print(Fore.GREEN + f"Ссылка для ключа '{key_to_delete}' удалена.")
            logging.info(f"Удалена ссылка с ключом: '{key_to_delete}'.")
        else:
            print(Fore.RED + f"Ключ '{key_to_delete}' не найден.")

    elif choice == 4:
        old_key = input("Введите текущий ключ: ")
        new_key = input("Введите новый ключ: ")
        if old_key in url_links:
            url_links[new_key] = url_links.pop(old_key)
            save_links(url_links)
            print(Fore.GREEN + f"Ключ '{old_key}' переименован в '{new_key}'.")
            logging.info(f"Ключ '{old_key}' переименован в '{new_key}'.")
        else:
            print(Fore.RED + f"Ключ '{old_key}' не найден.")

    elif choice == 5:
        show_available_keys(url_links)

    elif choice == 6:
        while True:
            print("\nНастройки:")
            print("1. Изменить пароль для входа")
            print("2. Пароль на открытие программы (сейчас: " + ("ВКЛ" if settings["password_required"] else "ВЫКЛ") + ")")
            print("3. Отображать ссылки (сейчас: " + ("ВКЛ" if settings["show_links"] else "ВЫКЛ") + ")")
            print("4. Уровень логирования (сейчас: " + settings.get("log_level", "INFO") + ")")
            print("5. Отладочные функции")
            print("6. Сброс программы")
            print("7. Назад")

            settings_choice = menu_option("Введите номер действия: ", range(1, 8))

            if settings_choice == 1:
                new_password = getpass.getpass("Введите новый пароль: ")
                settings["password"] = hash_password(new_password)
                save_settings(settings)
                print(Fore.GREEN + "Пароль изменен.")
            elif settings_choice == 2:
                settings["password_required"] = not settings["password_required"]
                save_settings(settings)
                status = "включено" if settings["password_required"] else "выключено"
                print(f"Требование пароля {status}.")
            elif settings_choice == 3:
                settings["show_links"] = not settings["show_links"]
                save_settings(settings)
                status = "включено" if settings["show_links"] else "выключено"
                print(f"Отображение ссылок {status}.")
            elif settings_choice == 4:
                print("\nВыберите уровень логирования:")
                print("1. DEBUG")
                print("2. INFO")
                print("3. WARNING")
                print("4. ERROR")
                print("5. CRITICAL")
                log_level_choice = menu_option("Введите номер уровня: ", range(1, 6))
                levels = ["DEBUG", "INFO", "WARNING", "ERROR", "CRITICAL"]
                settings["log_level"] = levels[log_level_choice - 1]
                save_settings(settings)
                set_log_level(settings)
            elif settings_choice == 5:
                run_debug_functions(url_links)
            elif settings_choice == 6:
                reset_program()
                url_links = load_links()
                settings = load_settings()
                set_log_level(settings)
                break
            elif settings_choice == 7:
                break

    elif choice == 7:
        count_links(url_links)

    elif choice == 8:
        filetypes = [("CSV files", "*.csv"),
                     ("JSON files", "*.json"),
                     ("YAML files", "*.yaml"),
                     ("XML files", "*.xml"),
                     ("DOCX files", "*.docx"),
                     ("TXT files", "*.txt"),
                     ("XLSX files", "*.xlsx"),
                     ("All files", "*.*")]
        filename = choose_file(save=True, filetypes=filetypes)
        if filename:
            format = filename.split('.')[-1].lower()
            export_links(url_links, filename, format)

    elif choice == 9:
        filetypes = [("CSV files", "*.csv"),
                     ("JSON files", "*.json"),
                     ("YAML files", "*.yaml"),
                     ("XML files", "*.xml"),
                     ("DOCX files", "*.docx"),
                     ("TXT files", "*.txt"),
                     ("XLSX files", "*.xlsx"),
                     ("All files", "*.*")]
        filename = choose_file(save=False, filetypes=filetypes)
        if filename:
            format = filename.split('.')[-1].lower()
            import_links(filename, format)

    elif choice == 10:
        query = input("Введите строку для поиска: ")
        print("Выберите тип поиска:")
        print("1. По ключевому слову")
        print("2. По регулярному выражению")
        search_type_choice = menu_option("Введите номер действия: ", [1, 2])
        search_type = 'keyword' if search_type_choice == 1 else 'regex'

        filters = {}
        filter_category = input("Фильтровать по категории (или оставьте пустым): ")
        if filter_category:
            filters['category'] = filter_category

        filter_date_from = input("Фильтровать по дате добавления (с) (YYYY-MM-DD, или оставьте пустым): ")
        if filter_date_from:
            filters['date_from'] = filter_date_from

        filter_date_to = input("Фильтровать по дате добавления (по) (YYYY-MM-DD, или оставьте пустым): ")
        if filter_date_to:
            filters['date_to'] = filter_date_to

        status_choice = input("Фильтровать по статусу ссылки (доступна/недоступна/нет): ").lower()
        if status_choice == 'доступна':
            filters['status'] = 'accessible'
        elif status_choice == 'недоступна':
            filters['status'] = 'inaccessible'

        search_links(url_links, query, search_type, filters)

    elif choice == 11:
        show_statistics(url_links)

    elif choice == 12:
        print(Fore.GREEN + "Выход из Link Manager.")
        logging.info("Программа завершена.")
        break

    else:
        print(Fore.RED + "Неверный ввод. Пожалуйста, попробуйте снова.")
