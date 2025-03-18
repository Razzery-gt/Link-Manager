
import webbrowser
import json
import os
import getpass
import bcrypt
import validators
import requests
import csv
import yaml
import xml.etree.ElementTree as ET
from docx import Document
import pandas as pd
import tkinter as tk
from tkinter import filedialog
import pyperclip
from datetime import datetime
from colorama import init, Fore

# Initialize colorama
init(autoreset=True)

# Constants
DOCUMENTS_DIR = os.path.join(os.path.expanduser("~"), "Documents", "LinkManager files")
LINKS_FILENAME = os.path.join(DOCUMENTS_DIR, 'url_links.json')
SETTINGS_FILENAME = os.path.join(DOCUMENTS_DIR, 'settings.json')

# Default settings
default_settings = {
    "password": bcrypt.hashpw("1234".encode(), bcrypt.gensalt()).decode(),
    "password_required": False,
    "show_links": True
}

# Default links
default_links = {
    "open_browser": {"url": "https://example.com", "date_added": str(datetime.now())},
    "google": {"url": "https://www.google.com", "date_added": str(datetime.now())},
    "yandex": {"url": "https://www.yandex.ru", "date_added": str(datetime.now())}
}

# Functions
def load_links():
    if os.path.exists(LINKS_FILENAME):
        try:
            with open(LINKS_FILENAME, 'r', encoding='utf-8') as f:
                links = json.load(f)
                # Convert old formats to the new format with date added
                for key, value in links.items():
                    if isinstance(value, str):  # Old format, URL only
                        links[key] = {"url": value, "date_added": str(datetime.now())}
                return links
        except (json.JSONDecodeError, IOError) as e:
            print(Fore.RED + f"Error loading links: {e}.")
    return default_links.copy()

def save_links(links):
    try:
        with open(LINKS_FILENAME, 'w', encoding='utf-8') as f:
            json.dump(links, f, ensure_ascii=False, indent=4)
    except IOError as e:
        print(Fore.RED + f"Error saving links: {e}.")

def load_settings():
    if os.path.exists(SETTINGS_FILENAME):
        try:
            with open(SETTINGS_FILENAME, 'r', encoding='utf-8') as f:
                return json.load(f)
        except (json.JSONDecodeError, IOError) as e:
            print(Fore.RED + f"Error loading settings: {e}.")
    return default_settings.copy()

def save_settings(settings):
    try:
        with open(SETTINGS_FILENAME, 'w', encoding='utf-8') as f:
            json.dump(settings, f, ensure_ascii=False, indent=4)
    except IOError as e:
        print(Fore.RED + f"Error saving settings: {e}.")

def open_browser(url):
    try:
        webbrowser.open(url)
        print(Fore.GREEN + f"Opening: {url}")
    except Exception as e:
        print(Fore.RED + f"Error opening browser: {e}")

def is_valid_url(url):
    return validators.url(url)

def check_url_accessibility(url):
    try:
        response = requests.get(url, timeout=5)
        return response.status_code == 200
    except requests.RequestException:
        return False

def show_available_keys(links):
    if links:
        print(Fore.CYAN + "Available keys to open in the browser:")
        for index, key in enumerate(links.keys(), 1):
            print(f"{index}. {Fore.YELLOW}{key} - {links[key]['url']} (Added: {links[key]['date_added']})")
    else:
        print(Fore.RED + "No available keys.")

def reset_program():
    if os.path.exists(LINKS_FILENAME):
        os.remove(LINKS_FILENAME)
    if os.path.exists(SETTINGS_FILENAME):
        os.remove(SETTINGS_FILENAME)
    print("Link Manager has been reset to default settings.")

def hash_password(password):
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def verify_password(stored_password, provided_password):
    return bcrypt.checkpw(provided_password.encode(), stored_password.encode())

def count_links(links):
    print(Fore.GREEN + f"Number of saved links: {len(links)}")

def copy_to_clipboard(text):
    try:
        pyperclip.copy(text)
        print(Fore.GREEN + "Link copied to clipboard.")
    except pyperclip.PyperclipException:
        print(Fore.RED + "Error copying to clipboard.  Make sure the pyperclip library is installed and the clipboard is working.")

def choose_file(save=False, filetypes=(("JSON files", "*.json"), ("All files", "*.*"))):
    root = tk.Tk()
    root.withdraw()  # Hide the main window
    if save:
        filepath = filedialog.asksaveasfilename(initialdir=DOCUMENTS_DIR, title="Choose a location to save", filetypes=filetypes)
    else:
        filepath = filedialog.askopenfilename(initialdir=DOCUMENTS_DIR, title="Choose a file to import", filetypes=filetypes)
    return filepath

# --- Export functions ---
def export_links(links, filename, format):
    try:
        if format == 'csv':
            export_to_csv(links, filename)
        elif format == 'json':
            export_to_json(links, filename)
        elif format == 'yaml':
            export_to_yaml(links, filename)
        elif format == 'xml':
            export_to_xml(links, filename)
        elif format == 'docx':
            export_to_docx(links, filename)
        elif format == 'txt':
            export_to_txt(links, filename)
        elif format == 'xlsx':
            export_to_xlsx(links, filename)
        else:
            print(Fore.RED + "Unsupported file format.")
            return

        print(Fore.GREEN + f"Links exported to {filename} in {format.upper()} format.")

    except Exception as e:
        print(Fore.RED + f"Error exporting links: {e}.")

def export_to_csv(links, filename):
    with open(filename, mode='w', newline='', encoding='utf-8') as csv_file:
        writer = csv.writer(csv_file)
        writer.writerow(['Key', 'URL', 'Date Added'])
        for key, data in links.items():
            writer.writerow([key, data['url'], data['date_added']])

def export_to_json(links, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        json.dump(links, f, ensure_ascii=False, indent=4)

def export_to_yaml(links, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        yaml.dump(links, f, allow_unicode=True, indent=4, stream=f)

def export_to_xml(links, filename):
    root = ET.Element("links")
    for key, data in links.items():
        link = ET.SubElement(root, "link")
        ET.SubElement(link, "key").text = key
        ET.SubElement(link, "url").text = data['url']
        ET.SubElement(link, "date_added").text = data['date_added']

    tree = ET.ElementTree(root)
    tree.write(filename, encoding='utf-8', xml_declaration=True)

def export_to_docx(links, filename):
    document = Document()
    document.add_heading('Links', level=1)
    for key, data in links.items():
        document.add_paragraph(f"Key: {key}")
        document.add_paragraph(f"URL: {data['url']}")
        document.add_paragraph(f"Date Added: {data['date_added']}")
        document.add_paragraph()

    document.save(filename)

def export_to_txt(links, filename):
    with open(filename, 'w', encoding='utf-8') as f:
        for key, data in links.items():
            f.write(f"Key: {key}\n")
            f.write(f"URL: {data['url']}\n")
            f.write(f"Date Added: {data['date_added']}\n\n")

def export_to_xlsx(links, filename):
    data = []
    for key, link_data in links.items():
        data.append([key, link_data['url'], link_data['date_added']])

    df = pd.DataFrame(data, columns=['Key', 'URL', 'Date Added'])
    df.to_excel(filename, index=False)

# --- Import functions ---
def import_links(filename, format):
    try:
        if format == 'csv':
            new_links = import_from_csv(filename)
        elif format == 'json':
            new_links = import_from_json(filename)
        elif format == 'yaml':
            new_links = import_from_yaml(filename)
        elif format == 'xml':
            new_links = import_from_xml(filename)
        elif format == 'docx':
            new_links = import_from_docx(filename)
        elif format == 'txt':
            new_links = import_from_txt(filename)
        elif format == 'xlsx':
            new_links = import_from_xlsx(filename)
        else:
            print(Fore.RED + "Unsupported file format.")
            return

        # Add new links, checking for duplicates
        for key, data in new_links.items():
            if key in url_links:
                print(Fore.YELLOW + f"Key '{key}' already exists. Skipping.")
            elif data['url'] in [link['url'] for link in url_links.values()]:
                print(Fore.YELLOW + f"Link '{data['url']}' already exists. Skipping.")
            else:
                url_links[key] = data
                print(Fore.GREEN + f"Imported link: {key} - {data['url']}")

        save_links(url_links)
        print(Fore.GREEN + f"Links imported from {filename} in {format.upper()} format.")

    except Exception as e:
        print(Fore.RED + f"Error importing links: {e}.")

def import_from_csv(filename):
    links = {}
    with open(filename, mode='r', encoding='utf-8') as csv_file:
        reader = csv.reader(csv_file)
        next(reader, None)  # Skip header
        for row in reader:
            if len(row) == 3:
                key, url, date_added = row
                if is_valid_url(url):
                    links[key] = {'url': url, 'date_added': date_added}
                else:
                    print(Fore.RED + f"Invalid URL '{url}' in row '{row}'. Skipping.")
    return links

def import_from_json(filename):
    with open(filename, 'r', encoding='utf-8') as f:
        return json.load(f)

def import_from_yaml(filename):
    with open(filename, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)

def import_from_xml(filename):
    tree = ET.parse(filename)
    root = tree.getroot()
    links = {}
    for link in root.findall('link'):
        key = link.find('key').text
        url = link.find('url').text
        date_added = link.find('date_added').text
        links[key] = {'url': url, 'date_added': date_added}
    return links

def import_from_docx(filename):
    document = Document(filename)
    links = {}
    key = None
    url = None
    date_added = None

    for paragraph in document.paragraphs:
        text = paragraph.text
        if text.startswith("Key: "):
            key = text[5:]
        elif text.startswith("URL: "):
            url = text[5:]
        elif text.startswith("Date Added: "):
            date_added = text[12:]

        if key and url and date_added:
            links[key] = {'url': url, 'date_added': date_added}
            key = None
            url = None
            date_added = None
    return links

def import_from_txt(filename):
    links = {}
    key = None
    url = None
    date_added = None

    with open(filename, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            if line.startswith("Key: "):
                key = line[5:]
            elif line.startswith("URL: "):
                url = line[5:]
            elif line.startswith("Date Added: "):
                date_added = line[12:]

            if key and url and date_added:
                links[key] = {'url': url, 'date_added': date_added}
                key = None
                url = None
                date_added = None
    return links

def import_from_xlsx(filename):
    df = pd.read_excel(filename)
    links = {}
    for index, row in df.iterrows():
        key = str(row['Key'])  # Convert key to string to handle numeric keys
        url = row['URL']
        date_added = str(row['Date Added'])  # Convert date to string
        links[key] = {'url': url, 'date_added': date_added}
    return links

def search_links(links, query):
    found_links = {key: data for key, data in links.items() if query.lower() in key.lower() or query.lower() in data['url'].lower()}
    if found_links:
        print(Fore.CYAN + "Found links:")
        for key, data in found_links.items():
            print(f"{Fore.YELLOW}{key}: {data['url']} (Added: {data['date_added']})")
    else:
        print(Fore.RED + "No links found.")

def menu_option(prompt, options):
    while True:
        try:
            choice = int(input(prompt))
            if choice in options:
                return choice
            else:
                print(Fore.RED + "Invalid input. Please try again.")
        except ValueError:
            print(Fore.RED + "Invalid input. Please try again.")

# Main loop
url_links = load_links()
settings = load_settings()

# Password check
if settings["password_required"]:
    password_attempts = 3
    while password_attempts > 0:
        password_input = getpass.getpass("Enter the password to access Link Manager: ")
        if verify_password(settings["password"], password_input):
            break
        else:
            password_attempts -= 1
            print(Fore.RED + f"Incorrect password. Attempts remaining: {password_attempts}")
    else:
        print(Fore.RED + "Access denied. Press any key to continue...")
        input()
        exit()

print(Fore.GREEN + "Welcome to Link Manager!")
print(Fore.GREEN + "Version: 1.0")

while True:
    print("\nChoose an action:")
    print("1. Open link by key")
    print("2. Add a new link or modify an existing one")
    print("3. Delete link")
    print("4. Rename link")
    print("5. Show available keys")
    print("6. Settings")
    print("7. Count saved links")
    print("8. Export links")
    print("9. Import links")
    print("10. Search links")
    print("11. Exit")

    choice = menu_option("Enter action number: ", range(1, 12))

    if choice == 1:
        if settings["show_links"]:
            show_available_keys(url_links)
        user_input = input("Enter the key to open the browser (or enter the number, or 'copy'): ").strip()

        if user_input.lower() == 'copy':
            key_to_copy = input("Enter the key of the link to copy: ")
            if key_to_copy in url_links:
                copy_to_clipboard(url_links[key_to_copy]['url'])
            else:
                print(Fore.RED + f"Key '{key_to_copy}' not found.")
        elif user_input.isdigit():
            selected_index = int(user_input) - 1
            if 0 <= selected_index < len(url_links):
                key_list = list(url_links.keys())
                selected_key = key_list[selected_index]
                selected_url = url_links[selected_key]['url']
                if check_url_accessibility(selected_url):
                    open_browser(selected_url)
                else:
                    print(Fore.RED + f"Failed to access URL: {selected_url}")
            else:
                print(Fore.RED + "Invalid key number.")
        elif user_input.lower() in [k.lower() for k in url_links.keys()]:
            selected_key = next(k for k in url_links.keys() if k.lower() == user_input.lower())
            selected_url = url_links[selected_key]['url']
            if check_url_accessibility(selected_url):
                open_browser(selected_url)
            else:
                print(Fore.RED + f"Failed to access URL: {selected_url}")
        else:
            print(Fore.RED + f"Key '{user_input}' not found.")

    elif choice == 2:
        new_key = input("Enter key: ")
        new_url = input("Enter URL: ")

        if not is_valid_url(new_url):
            print(Fore.RED + "Invalid URL. Please enter a valid URL.")
            continue

        if new_url in [link['url'] for link in url_links.values()]:
            print(Fore.RED + "This link already exists.")
            continue

        url_links[new_key] = {"url": new_url, "date_added": str(datetime.now())}
        save_links(url_links)
        print(Fore.GREEN + f"Link for key '{new_key}' added/updated.")

    elif choice == 3:
        key_to_delete = input("Enter key to delete: ")
        if key_to_delete in url_links:
            del url_links[key_to_delete]
            save_links(url_links)
            print(Fore.GREEN + f"Link for key '{key_to_delete}' deleted.")
        else:
            print(Fore.RED + f"Key '{key_to_delete}' not found.")

    elif choice == 4:
        old_key = input("Enter current key: ")
        new_key = input("Enter new key: ")
        if old_key in url_links:
            url_links[new_key] = url_links.pop(old_key)
            save_links(url_links)
            print(Fore.GREEN + f"Key '{old_key}' renamed to '{new_key}'.")
        else:
            print(Fore.RED + f"Key '{old_key}' not found.")

    elif choice == 5:
        show_available_keys(url_links)

    elif choice == 6:
        while True:
            print("\nSettings:")
            print("1. Change login password")
            print("2. Password on program launch (currently: " + ("ON" if settings["password_required"] else "OFF") + ")")
            print("3. Display links (currently: " + ("ON" if settings["show_links"] else "OFF") + ")")
            print("4. Reset program")
            print("5. Back")

            settings_choice = menu_option("Enter action number: ", range(1, 6))

            if settings_choice == 1:
                new_password = getpass.getpass("Enter new password: ")
                settings["password"] = hash_password(new_password)
                save_settings(settings)
                print(Fore.GREEN + "Password changed.")
            elif settings_choice == 2:
                settings["password_required"] = not settings["password_required"]
                save_settings(settings)
                status = "enabled" if settings["password_required"] else "disabled"
                print(f"Password requirement {status}.")
            elif settings_choice == 3:
                settings["show_links"] = not settings["show_links"]
                save_settings(settings)
                status = "enabled" if settings["show_links"] else "disabled"
                print(f"Displaying links {status}.")
            elif settings_choice == 4:
                reset_program()
                url_links = load_links()
                settings = load_settings()
                break
            elif settings_choice == 5:
                break

    elif choice == 7:
        count_links(url_links)

    elif choice == 8:
        filetypes = [("CSV files", "*.csv"),
                     ("JSON files", "*.json"),
                     ("YAML files", "*.yaml"),
                     ("XML files", "*.xml"),
                     ("DOCX files", "*.docx"),
                     ("TXT files", "*.txt"),
                     ("XLSX files", "*.xlsx"),
                     ("All files", "*.*")]
        filename = choose_file(save=True, filetypes=filetypes)
        if filename:
            format = filename.split('.')[-1].lower()
            export_links(url_links, filename, format)

    elif choice == 9:
        filetypes = [("CSV files", "*.csv"),
                     ("JSON files", "*.json"),
                     ("YAML files", "*.yaml"),
                     ("XML files", "*.xml"),
                     ("DOCX files", "*.docx"),
                     ("TXT files", "*.txt"),
                     ("XLSX files", "*.xlsx"),
                     ("All files", "*.*")]
        filename = choose_file(save=False, filetypes=filetypes)
        if filename:
            format = filename.split('.')[-1].lower()
            import_links(filename, format)

    elif choice == 10:
        query = input("Enter keyword to search: ")
        search_links(url_links, query)

    elif choice == 11:
        print(Fore.GREEN + "Exiting Link Manager.")
        break

    else:
        print(Fore.RED + "Invalid input. Please try again.")
